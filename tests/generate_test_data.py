from docx import Document
from fpdf import FPDF
import os

def create_test_docx(filename):
    doc = Document()
    doc.add_heading('Wo ist meine Doku - Testdokument (Office)', 0)
    doc.add_paragraph('Dies ist ein Testdokument für die automatisierte Dokumentenextraktion.')
    doc.add_paragraph('Hier ist ein komplexes deutsches Wort: Donaudampfschifffahrtsgesellschaftskapitän.')
    doc.add_paragraph('Das System sollte in der Lage sein, dieses Wort in seine Bestandteile zu zerlegen.')
    doc.save(filename)
    print(f"Created {filename}")

def create_test_pdf(filename):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, text="Wo ist meine Doku - Testdokument (PDF)", border=1, align='C')
    pdf.ln(10)
    pdf.multi_cell(0, 10, text="Dies ist ein PDF-Testdokument.\n\n"
                            "Wir testen die Extraktion von Text aus komplexen Layouts.\n"
                            "Wichtiges Wort: Rindfleischetikettierungsüberwachungsaufgabenübertragungsgesetz.\n"
                            "Umlaute Test: ÄÖÜäß.")
    pdf.output(filename)
    print(f"Created {filename}")

if __name__ == "__main__":
    os.makedirs('tests/data', exist_ok=True)
    create_test_docx('tests/data/test_office.docx')
    create_test_pdf('tests/data/test_pdf.pdf')
